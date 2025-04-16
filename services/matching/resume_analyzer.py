"""
Analisador de currículos para extração de informações estruturadas.

Este módulo contém a implementação do analisador de currículos que extrai
informações estruturadas de currículos em diversos formatos.
"""

import logging
import re
import os
import json
from typing import Dict, List, Optional, Any, Union
from datetime import datetime
import tempfile

# Bibliotecas para processamento de documentos
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None

try:
    from pdfminer.high_level import extract_text as pdf_extract_text
except ImportError:
    pdf_extract_text = None

# NLP
try:
    import nltk
    from nltk.tokenize import word_tokenize, sent_tokenize
    from nltk.corpus import stopwords

    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
except ImportError:
    nltk = None

from ..ai.claude_service import ClaudeService

logger = logging.getLogger(__name__)


class ResumeAnalyzer:
    """
    Analisador de currículos para extração de informações estruturadas.

    Esta classe implementa métodos para extrair dados de currículos em
    diversos formatos, como PDF, DOCX e texto plano.
    """

    def __init__(self, claude_service: ClaudeService = None):
        """
        Inicializa o analisador de currículos.

        Args:
            claude_service: Serviço de integração com Claude 3.7 Sonnet
        """
        self.claude_service = claude_service

        # Expressões regulares para extração de informações
        self.regex_patterns = {
            'email': r'[\w\.-]+@[\w\.-]+\.\w+',
            'phone': r'(?:\+\d{1,3})?(?:\(\d{2,3}\)|\d{2,3})[-.\s]?\d{3,5}[-.\s]?\d{4}',
            'url': r'(?:https?://)?(?:www\.)?[\w-]+\.[\w.-]+(?:/[\w-]*)*',
            'education': r'(?:formation|education|formação|educação|acadêmico|academic)',
            'experience': r'(?:experience|experiência|professional|profissional|work|trabalho)',
            'skills': r'(?:skills|habilidades|competências|competencies|conhecimentos|knowledge)',
            'languages': r'(?:languages|idiomas|línguas)'
        }

        # Listas de habilidades técnicas e não-técnicas comuns
        self.technical_skills = [
            'python', 'java', 'javascript', 'typescript', 'c#', 'c\+\+', 'php', 'ruby', 'swift',
            'kotlin', 'go', 'rust', 'html', 'css', 'sql', 'nosql', 'react', 'angular', 'vue',
            'node', 'django', 'flask', 'spring', 'asp\.net', 'laravel', 'rails', 'express',
            'mongodb', 'postgresql', 'mysql', 'oracle', 'sqlite', 'redis', 'aws', 'azure',
            'gcp', 'docker', 'kubernetes', 'jenkins', 'ci/cd', 'git', 'github', 'gitlab',
            'jira', 'scrum', 'kanban', 'agile', 'machine learning', 'ml', 'ai', 'deep learning',
            'data science', 'big data', 'hadoop', 'spark', 'tableau', 'power bi', 'excel',
            'word', 'powerpoint', 'photoshop', 'illustrator', 'figma', 'sketch'
        ]

        self.soft_skills = [
            'communication', 'teamwork', 'leadership', 'problem solving', 'critical thinking',
            'decision making', 'time management', 'organization', 'adaptability', 'flexibility',
            'creativity', 'emotional intelligence', 'conflict resolution', 'negotiation',
            'persuasion', 'collaboration', 'interpersonal', 'trabalho em equipe', 'liderança',
            'resolução de problemas', 'pensamento crítico', 'tomada de decisão', 'gestão de tempo',
            'organização', 'adaptabilidade', 'flexibilidade', 'criatividade', 'inteligência emocional',
            'resolução de conflitos', 'negociação', 'persuasão', 'colaboração', 'interpessoal',
            'comunicação', 'proatividade', 'empatia', 'resiliência'
        ]

    def analyze(self, file_data: bytes, file_name: str, use_claude: bool = True) -> Dict[str, Any]:
        """
        Analisa um arquivo de currículo e extrai informações estruturadas.

        Args:
            file_data: Conteúdo binário do arquivo
            file_name: Nome do arquivo com extensão
            use_claude: Se True, usa o Claude 3.7 para análise

        Returns:
            Dicionário com informações estruturadas do currículo
        """
        logger.info(f"Analisando currículo: {file_name}")

        # Extrai o texto do currículo
        text = self.extract_text(file_data, file_name)

        if not text:
            logger.error(f"Não foi possível extrair texto do arquivo: {file_name}")
            return {
                'error': "Não foi possível extrair texto do arquivo",
                'extracted_text': None
            }

        # Se utilizarmos o Claude para análise, delegamos a extração
        if use_claude and self.claude_service:
            try:
                resume_data = self.claude_service.analyze_resume(text)

                # Adiciona o texto extraído no resultado
                resume_data['extracted_text'] = text

                logger.info(f"Currículo analisado pelo Claude: {file_name}")
                return resume_data

            except Exception as e:
                logger.error(f"Erro ao analisar currículo com Claude: {str(e)}")
                logger.info("Recorrendo ao algoritmo de análise interno")
                # Continua com o algoritmo interno em caso de falha

        # Caso contrário, usamos o algoritmo interno
        try:
            resume_data = self._extract_resume_data(text)

            # Adiciona o texto extraído no resultado
            resume_data['extracted_text'] = text

            logger.info(f"Currículo analisado internamente: {file_name}")
            return resume_data

        except Exception as e:
            logger.error(f"Erro ao analisar currículo internamente: {str(e)}")

            # Retorna apenas o texto extraído em caso de erro
            return {
                'error': str(e),
                'extracted_text': text
            }

    def extract_text(self, file_data: bytes, file_name: str) -> Optional[str]:
        """
        Extrai texto de um arquivo de currículo.

        Args:
            file_data: Conteúdo binário do arquivo
            file_name: Nome do arquivo com extensão

        Returns:
            Texto extraído do arquivo ou None em caso de erro
        """
        # Identifica a extensão do arquivo
        extension = os.path.splitext(file_name)[1].lower()

        try:
            # Cria um arquivo temporário para processamento
            with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as temp_file:
                temp_file.write(file_data)
                temp_file_path = temp_file.name

            try:
                # Extrai o texto com base na extensão
                if extension == '.pdf':
                    text = self._extract_text_from_pdf(temp_file_path)
                elif extension in ['.docx', '.doc']:
                    text = self._extract_text_from_docx(temp_file_path)
                elif extension == '.txt':
                    with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read()
                else:
                    logger.warning(f"Formato de arquivo não suportado: {extension}")
                    text = None

                return text

            finally:
                # Remove o arquivo temporário
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)

        except Exception as e:
            logger.error(f"Erro ao extrair texto do arquivo {file_name}: {str(e)}")
            return None

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extrai texto de um arquivo PDF.

        Args:
            file_path: Caminho para o arquivo PDF

        Returns:
            Texto extraído do PDF
        """
        # Tenta com pdfminer primeiro (melhor para textos multilínguas)
        if pdf_extract_text:
            try:
                text = pdf_extract_text(file_path)
                if text and len(text) > 100:
                    return text
            except Exception as e:
                logger.warning(f"Erro ao extrair texto com pdfminer: {str(e)}")

        # Recorre ao PyPDF2 se pdfminer falhar
        if PyPDF2:
            try:
                text = ""
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text += page.extract_text() + "\n"
                return text
            except Exception as e:
                logger.error(f"Erro ao extrair texto com PyPDF2: {str(e)}")
                raise
        else:
            raise ImportError("Nenhuma biblioteca de processamento de PDF disponível")

    def _extract_text_from_docx(self, file_path: str) -> str:
        """
        Extrai texto de um arquivo DOCX.

        Args:
            file_path: Caminho para o arquivo DOCX

        Returns:
            Texto extraído do DOCX
        """
        if not docx:
            raise ImportError("Biblioteca python-docx não está disponível")

        try:
            doc = docx.Document(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Extrai texto de tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                text += "\n"

            return text
        except Exception as e:
            logger.error(f"Erro ao extrair texto do DOCX: {str(e)}")
            raise

    def _extract_resume_data(self, text: str) -> Dict[str, Any]:
        """
        Extrai informações estruturadas do texto do currículo.

        Args:
            text: Texto extraído do currículo

        Returns:
            Dicionário com informações estruturadas do currículo
        """
        # Normaliza o texto
        text = self._normalize_text(text)

        # Divide o texto em seções
        sections = self._split_into_sections(text)

        # Extrai informações pessoais
        personal_info = self._extract_personal_info(text)

        # Extrai formação acadêmica
        education = self._extract_education(sections.get('education', text))

        # Extrai experiência profissional
        experience = self._extract_experience(sections.get('experience', text))

        # Extrai habilidades
        skills = self._extract_skills(sections.get('skills', text))

        # Extrai idiomas
        languages = self._extract_languages(sections.get('languages', text))

        # Constrói o resultado
        result = {
            'personal_info': personal_info,
            'education': education,
            'experience': experience,
            'skills': skills,
            'languages': languages,
            'analyzed_at': datetime.utcnow().isoformat()
        }

        return result

    def _normalize_text(self, text: str) -> str:
        """
        Normaliza o texto para facilitar o processamento.

        Args:
            text: Texto original

        Returns:
            Texto normalizado
        """
        # Remove caracteres de controle
        text = ''.join(ch if ord(ch) < 128 or ord(ch) > 160 else ' ' for ch in text)

        # Normaliza quebras de linha
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Normaliza espaços
        text = re.sub(r' {2,}', ' ', text)

        return text.strip()

    def _split_into_sections(self, text: str) -> Dict[str, str]:
        """
        Divide o texto do currículo em seções.

        Args:
            text: Texto normalizado do currículo

        Returns:
            Dicionário com o texto de cada seção
        """
        sections = {}

        # Divide o texto em linhas
        lines = text.split('\n')

        current_section = None
        section_content = []

        for line in lines:
            line = line.strip()

            # Pula linhas vazias
            if not line:
                if current_section:
                    section_content.append("")
                continue

            # Verifica se a linha é um cabeçalho de seção
            if self._is_section_header(line):
                # Salva a seção anterior, se houver
                if current_section:
                    sections[current_section] = '\n'.join(section_content)

                # Identifica o tipo da nova seção
                section_type = self._identify_section_type(line)
                current_section = section_type
                section_content = []
            else:
                # Adiciona a linha à seção atual
                if current_section:
                    section_content.append(line)
                else:
                    # Caso ainda não tenha identificado uma seção, considera como informações pessoais
                    if 'personal_info' not in sections:
                        sections['personal_info'] = ""
                    sections['personal_info'] += line + "\n"

        # Salva a última seção
        if current_section:
            sections[current_section] = '\n'.join(section_content)

        return sections

    def _is_section_header(self, line: str) -> bool:
        """
        Verifica se uma linha é um cabeçalho de seção.

        Args:
            line: Linha de texto

        Returns:
            True se a linha for um cabeçalho de seção, False caso contrário
        """
        # Cabeçalhos geralmente são curtos
        if len(line) > 50:
            return False

        # Cabeçalhos geralmente estão em maiúsculas ou com a primeira letra maiúscula
        if line.isupper() or line.istitle():
            return True

        # Cabeçalhos geralmente terminam com dois pontos
        if line.endswith(':'):
            return True

        # Cabeçalhos comuns
        common_headers = [
            'education', 'formation', 'academic', 'formação', 'educação', 'acadêmico',
            'experience', 'professional', 'work', 'experiência', 'profissional', 'trabalho',
            'skills', 'competencies', 'habilidades', 'competências', 'conhecimentos',
            'languages', 'idiomas', 'línguas', 'personal', 'profile', 'perfil', 'pessoal',
            'objective', 'objetivo', 'summary', 'resumo', 'qualifications', 'qualificações',
            'certifications', 'certificações', 'projects', 'projetos', 'awards', 'prêmios'
        ]

        line_lower = line.lower()
        for header in common_headers:
            if header in line_lower and len(line) < 30:
                return True

        return False

    def _identify_section_type(self, header: str) -> str:
        """
        Identifica o tipo de seção com base no cabeçalho.

        Args:
            header: Texto do cabeçalho

        Returns:
            Tipo da seção: education, experience, skills, languages ou other
        """
        header_lower = header.lower()

        # Seção de educação
        if re.search(self.regex_patterns['education'], header_lower):
            return 'education'

        # Seção de experiência
        if re.search(self.regex_patterns['experience'], header_lower):
            return 'experience'

        # Seção de habilidades
        if re.search(self.regex_patterns['skills'], header_lower):
            return 'skills'

        # Seção de idiomas
        if re.search(self.regex_patterns['languages'], header_lower):
            return 'languages'

        # Se não for identificada, retorna "other"
        return 'other'

    def _extract_personal_info(self, text: str) -> Dict[str, Any]:
        """
        Extrai informações pessoais do texto.

        Args:
            text: Texto do currículo

        Returns:
            Dicionário com informações pessoais
        """
        personal_info = {}

        # Extrai o nome (assumindo que está no início do documento)
        lines = text.split('\n')
        if lines:
            # O nome geralmente é a primeira linha não vazia
            for line in lines:
                if line.strip():
                    # Verifica se a linha parece um nome (não um email, telefone, etc.)
                    if not re.search(self.regex_patterns['email'], line) and \
                            not re.search(self.regex_patterns['phone'], line) and \
                            not re.search(self.regex_patterns['url'], line) and \
                            len(line) < 50:
                        personal_info['name'] = line.strip()
                        break

        # Extrai email
        email_match = re.search(self.regex_patterns['email'], text)
        if email_match:
            personal_info['email'] = email_match.group(0)

        # Extrai telefone
        phone_match = re.search(self.regex_patterns['phone'], text)
        if phone_match:
            personal_info['phone'] = phone_match.group(0)

        # Extrai URLs (LinkedIn, GitHub, etc.)
        url_matches = re.finditer(self.regex_patterns['url'], text)
        urls = []
        for match in url_matches:
            url = match.group(0)
            if url not in urls:
                urls.append(url)

        if urls:
            personal_info['urls'] = urls

        return personal_info

    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """
        Extrai informações sobre formação acadêmica.

        Args:
            text: Texto da seção de educação

        Returns:
            Lista de formações acadêmicas
        """
        if not text:
            return []

        education_list = []

        # Divide o texto em parágrafos ou blocos
        paragraphs = re.split(r'\n\s*\n', text)

        for paragraph in paragraphs:
            if not paragraph.strip():
                continue

            education_entry = {}

            # Extrai o grau acadêmico
            degree_patterns = [
                r'(?:Bacharelado|Bacharel|Licenciatura|Tecnólogo|Técnico|MBA|Mestrado|Doutorado|Pós-graduação|Especialização)',
                r'(?:Bachelor|Master|PhD|Doctorate|Post-graduate|Graduate|Undergraduate|Technical)'
            ]

            for pattern in degree_patterns:
                degree_match = re.search(pattern, paragraph, re.IGNORECASE)
                if degree_match:
                    education_entry['degree'] = degree_match.group(0)
                    break

            # Extrai a instituição
            institution_patterns = [
                r'(?:Universidade|Faculdade|Instituto|Escola|Colégio)\s+[\w\s]+',
                r'(?:University|College|Institute|School)\s+[\w\s]+'
            ]

            for pattern in institution_patterns:
                institution_match = re.search(pattern, paragraph, re.IGNORECASE)
                if institution_match:
                    education_entry['institution'] = institution_match.group(0)
                    break

            # Extrai a área de estudo
            if 'degree' in education_entry:
                # Procura pela área após o grau
                degree_pos = paragraph.find(education_entry['degree']) + len(education_entry['degree'])
                remaining_text = paragraph[degree_pos:].strip()

                # Procura por padrões como "em X", "in X"
                study_field_match = re.search(r'(?:em|in|of)\s+([\w\s]+)(?:,|\.|$)', remaining_text, re.IGNORECASE)
                if study_field_match:
                    education_entry['field_of_study'] = study_field_match.group(1).strip()

            # Extrai datas
            date_pattern = r'(?:Jan|Fev|Mar|Abr|Mai|Jun|Jul|Ago|Set|Out|Nov|Dez|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[\s\./]+\d{4}'
            date_pattern_alt = r'\d{1,2}[\s\./]+\d{1,2}[\s\./]+\d{4}'
            date_pattern_year = r'\d{4}\s*(?:-|–|a|to)\s*(?:presente|present|\d{4})'

            dates = []
            for pattern in [date_pattern, date_pattern_alt, date_pattern_year]:
                date_matches = re.finditer(pattern, paragraph, re.IGNORECASE)
                for match in date_matches:
                    dates.append(match.group(0))

            if dates:
                if len(dates) >= 2:
                    education_entry['start_date'] = dates[0]
                    education_entry['end_date'] = dates[1]
                else:
                    education_entry['date'] = dates[0]

            # Adiciona a entrada se tiver pelo menos instituição ou grau
            if 'institution' in education_entry or 'degree' in education_entry:
                education_list.append(education_entry)

        return education_list

    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """
        Extrai informações sobre experiência profissional.

        Args:
            text: Texto da seção de experiência

        Returns:
            Lista de experiências profissionais
        """
        if not text:
            return []

        experience_list = []

        # Divide o texto em parágrafos ou blocos
        paragraphs = re.split(r'\n\s*\n', text)

        for paragraph in paragraphs:
            if not paragraph.strip():
                continue

            experience_entry = {}
            lines = paragraph.split('\n')

            # A primeira linha geralmente contém o cargo
            if lines:
                experience_entry['title'] = lines[0].strip()

                # Se o cargo estiver junto com a empresa, tenta separá-los
                if ' - ' in experience_entry['title'] or ' at ' in experience_entry['title'] or ' na ' in \
                        experience_entry['title']:
                    parts = re.split(r'\s+(?:-|at|na|no)\s+', experience_entry['title'], maxsplit=1)
                    if len(parts) == 2:
                        experience_entry['title'] = parts[0].strip()
                        experience_entry['company'] = parts[1].strip()

            # A segunda linha geralmente contém a empresa se não estiver na primeira
            if len(lines) > 1 and 'company' not in experience_entry:
                experience_entry['company'] = lines[1].strip()

            # Extrai datas
            date_pattern = r'(?:Jan|Fev|Mar|Abr|Mai|Jun|Jul|Ago|Set|Out|Nov|Dez|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[\s\./]+\d{4}'
            date_pattern_alt = r'\d{1,2}[\s\./]+\d{1,2}[\s\./]+\d{4}'
            date_pattern_year = r'\d{4}\s*(?:-|–|a|to)\s*(?:presente|present|\d{4}|atual|current)'

            dates = []
            for pattern in [date_pattern, date_pattern_alt, date_pattern_year]:
                date_matches = re.finditer(pattern, paragraph, re.IGNORECASE)
                for match in date_matches:
                    dates.append(match.group(0))

            if dates:
                if len(dates) >= 2:
                    experience_entry['start_date'] = dates[0]
                    experience_entry['end_date'] = dates[1]
                else:
                    # Se houver apenas uma data e tiver "presente" ou "atual", é a data de início
                    if 'presente' in paragraph.lower() or 'present' in paragraph.lower() or 'atual' in paragraph.lower() or 'current' in paragraph.lower():
                        experience_entry['start_date'] = dates[0]
                        experience_entry['end_date'] = 'Presente'
                    else:
                        experience_entry['date'] = dates[0]

            # Extrai descrição/responsabilidades
            description_lines = []
            description_started = False

            for i, line in enumerate(lines):
                if i <= 1:  # Pula as primeiras linhas (título e empresa)
                    continue

                # Verifica se a linha contém datas
                has_date = False
                for pattern in [date_pattern, date_pattern_alt, date_pattern_year]:
                    if re.search(pattern, line, re.IGNORECASE):
                        has_date = True
                        break

                if has_date and not description_started:
                    continue

                description_started = True
                description_lines.append(line.strip())

            if description_lines:
                experience_entry['description'] = '\n'.join(description_lines)

            # Adiciona a entrada se tiver pelo menos título ou empresa
            if 'title' in experience_entry or 'company' in experience_entry:
                experience_list.append(experience_entry)

        return experience_list

    def _extract_skills(self, text: str) -> Dict[str, List[str]]:
        """
        Extrai habilidades do texto.

        Args:
            text: Texto da seção de habilidades

        Returns:
            Dicionário com habilidades técnicas e não-técnicas
        """
        if not text:
            return {'technical': [], 'soft': []}

        skills = {'technical': [], 'soft': []}

        # Extrai todas as linhas que podem conter habilidades
        lines = re.split(r'[,\n•\-]', text)

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Verifica se a linha menciona uma habilidade técnica
            for skill in self.technical_skills:
                if re.search(r'\b' + skill + r'\b', line.lower()):
                    # Extrai a habilidade completa
                    if len(line) < 50:  # Evita linhas muito longas
                        if line not in skills['technical']:
                            skills['technical'].append(line)
                    else:
                        # Se a linha for muito longa, extrai apenas a habilidade
                        match = re.search(r'\b' + skill + r'(?:\s+\w+){0,2}\b', line.lower())
                        if match and match.group(0) not in skills['technical']:
                            skills['technical'].append(match.group(0).capitalize())
                    break

            # Verifica se a linha menciona uma habilidade não-técnica
            for skill in self.soft_skills:
                if re.search(r'\b' + skill + r'\b', line.lower()):
                    # Extrai a habilidade completa
                    if len(line) < 50:  # Evita linhas muito longas
                        if line not in skills['soft']:
                            skills['soft'].append(line)
                    else:
                        # Se a linha for muito longa, extrai apenas a habilidade
                        match = re.search(r'\b' + skill + r'(?:\s+\w+){0,2}\b', line.lower())
                        if match and match.group(0) not in skills['soft']:
                            skills['soft'].append(match.group(0).capitalize())
                    break

        # Se não encontrou habilidades específicas, procura no texto completo
        if not skills['technical'] and not skills['soft']:
            for skill in self.technical_skills:
                if re.search(r'\b' + skill + r'\b', text.lower()):
                    skills['technical'].append(skill.capitalize())

            for skill in self.soft_skills:
                if re.search(r'\b' + skill + r'\b', text.lower()):
                    skills['soft'].append(skill.capitalize())

        return skills

    def _extract_languages(self, text: str) -> List[Dict[str, str]]:
        """
        Extrai idiomas do texto.

        Args:
            text: Texto da seção de idiomas

        Returns:
            Lista de idiomas com nível de proficiência
        """
        if not text:
            return []

        languages = []

        # Lista de idiomas comuns
        common_languages = [
            'português', 'portuguese', 'inglês', 'english', 'espanhol', 'spanish',
            'francês', 'french', 'alemão', 'german', 'italiano', 'italian',
            'japonês', 'japanese', 'chinês', 'chinese', 'mandarim', 'mandarin',
            'russo', 'russian', 'árabe', 'arabic', 'hindi', 'holandês', 'dutch'
        ]

        # Lista de níveis de proficiência
        proficiency_levels = [
            'nativo', 'native', 'fluente', 'fluent', 'avançado', 'advanced',
            'intermediário', 'intermediate', 'básico', 'basic', 'elementar', 'elementary',
            'proficiente', 'proficient', 'limitado', 'limited', 'conversacional',
            'conversational', 'técnico', 'technical', 'a1', 'a2', 'b1', 'b2', 'c1', 'c2'
        ]

        # Divide o texto em linhas
        lines = text.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Procura por idiomas na linha
            language_found = None
            for language in common_languages:
                if language.lower() in line.lower():
                    language_found = language.capitalize()
                    break

            if not language_found:
                continue

            # Procura por nível de proficiência
            proficiency = None
            for level in proficiency_levels:
                if level.lower() in line.lower():
                    proficiency = level.capitalize()
                    break

            languages.append({
                'language': language_found,
                'proficiency': proficiency or 'Não especificado'
            })

        return languages