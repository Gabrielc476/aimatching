"""
Text extraction utilities for LinkedIn Job Matcher.
"""

import io
import logging
from typing import Optional, Dict, Any, Tuple
import os
import tempfile

logger = logging.getLogger(__name__)


def extract_text_from_pdf(pdf_content: bytes) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text from a PDF document.

    Args:
        pdf_content: Binary PDF content

    Returns:
        Tuple of (extracted_text, metadata)
    """
    try:
        import PyPDF2
        from pdfminer.high_level import extract_text as pdfminer_extract

        text = ""
        metadata = {}

        # Extract text with PyPDF2
        with io.BytesIO(pdf_content) as pdf_stream:
            try:
                reader = PyPDF2.PdfReader(pdf_stream)

                # Extract metadata
                if reader.metadata:
                    for key, value in reader.metadata.items():
                        # Remove the leading slash in metadata keys if present
                        clean_key = key[1:] if key.startswith('/') else key
                        metadata[clean_key] = value

                # Extract text from each page
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            except Exception as e:
                logger.warning(f"PyPDF2 extraction failed, falling back to pdfminer: {str(e)}")

                # If PyPDF2 fails, try with pdfminer
                pdf_stream.seek(0)
                text = pdfminer_extract(pdf_stream)

        # If both methods failed to extract meaningful text, try using an OCR approach
        if not text or len(text.strip()) < 50:
            logger.warning("Text extraction yielded little text, attempting OCR")
            text, ocr_metadata = extract_text_with_ocr(pdf_content)
            metadata.update(ocr_metadata)

        return text, metadata
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        return "", {"error": str(e)}


def extract_text_from_docx(docx_content: bytes) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text from a DOCX document.

    Args:
        docx_content: Binary DOCX content

    Returns:
        Tuple of (extracted_text, metadata)
    """
    try:
        import docx

        text = ""
        metadata = {}

        # Create a temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_filename = temp_file.name
            temp_file.write(docx_content)

        try:
            # Open the document
            doc = docx.Document(temp_filename)

            # Extract document properties
            try:
                core_properties = doc.core_properties
                if core_properties:
                    metadata = {
                        "author": core_properties.author,
                        "created": core_properties.created,
                        "modified": core_properties.modified,
                        "title": core_properties.title,
                        "subject": core_properties.subject,
                        "keywords": core_properties.keywords,
                        "category": core_properties.category,
                        "comments": core_properties.comments
                    }
                    # Remove None values
                    metadata = {k: v for k, v in metadata.items() if v is not None}
            except Exception as e:
                logger.warning(f"Failed to extract DOCX metadata: {str(e)}")

            # Extract text from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"

            return text, metadata
        finally:
            # Clean up the temporary file
            try:
                os.unlink(temp_filename)
            except Exception:
                pass
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        return "", {"error": str(e)}


def extract_text_with_ocr(content: bytes) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text from an image or PDF using OCR.

    Args:
        content: Binary content

    Returns:
        Tuple of (extracted_text, metadata)
    """
    try:
        import pytesseract
        from PIL import Image
        from pdf2image import convert_from_bytes

        metadata = {"method": "ocr"}
        text = ""

        # Convert PDF to images
        images = convert_from_bytes(content)

        # Process each image with OCR
        for i, image in enumerate(images):
            page_text = pytesseract.image_to_string(image, lang='por+eng')
            if page_text:
                text += page_text + "\n\n"

        return text, metadata
    except Exception as e:
        logger.error(f"Error extracting text with OCR: {str(e)}")
        return "", {"error": str(e), "method": "ocr"}


def extract_text_from_txt(txt_content: bytes) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text from a plain text file.

    Args:
        txt_content: Binary text content

    Returns:
        Tuple of (extracted_text, metadata)
    """
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'windows-1252']
        text = ""
        error = None

        for encoding in encodings:
            try:
                text = txt_content.decode(encoding)
                error = None
                break
            except UnicodeDecodeError as e:
                error = e

        if not text and error:
            logger.error(f"Failed to decode text file with available encodings: {str(error)}")
            return "", {"error": "Encoding detection failed"}

        return text, {"format": "txt"}
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {str(e)}")
        return "", {"error": str(e)}


def extract_text_from_rtf(rtf_content: bytes) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text from an RTF document.

    Args:
        rtf_content: Binary RTF content

    Returns:
        Tuple of (extracted_text, metadata)
    """
    try:
        import striprtf.striprtf

        # Convert binary data to string
        rtf_string = rtf_content.decode('latin-1', errors='ignore')

        # Strip RTF formatting
        text = striprtf.striprtf.rtf_to_text(rtf_string)

        return text, {"format": "rtf"}
    except Exception as e:
        logger.error(f"Error extracting text from RTF: {str(e)}")
        return "", {"error": str(e)}


def extract_text_from_odt(odt_content: bytes) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text from an ODT document.

    Args:
        odt_content: Binary ODT content

    Returns:
        Tuple of (extracted_text, metadata)
    """
    try:
        import zipfile
        import xml.etree.ElementTree as ET

        metadata = {}
        text = ""

        # ODT files are actually ZIP files
        with zipfile.ZipFile(io.BytesIO(odt_content)) as odt_zip:
            # Extract content.xml
            content_xml = odt_zip.read('content.xml')

            # Parse XML
            root = ET.fromstring(content_xml)

            # Define namespace
            ns = {'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0'}

            # Extract text from paragraph nodes
            for paragraph in root.findall('.//text:p', ns):
                if paragraph.text:
                    text += paragraph.text + "\n"

                # Extract text from span nodes within paragraphs
                for child in paragraph:
                    if child.tag.endswith('}span') and child.text:
                        text += child.text

                text += "\n"

            # Try to extract metadata
            try:
                meta_xml = odt_zip.read('meta.xml')
                meta_root = ET.fromstring(meta_xml)

                meta_ns = {
                    'office': 'urn:oasis:names:tc:opendocument:xmlns:office:1.0',
                    'dc': 'http://purl.org/dc/elements/1.1/',
                    'meta': 'urn:oasis:names:tc:opendocument:xmlns:meta:1.0'
                }

                title = meta_root.find('.//dc:title', meta_ns)
                if title is not None and title.text:
                    metadata['title'] = title.text

                creator = meta_root.find('.//dc:creator', meta_ns)
                if creator is not None and creator.text:
                    metadata['author'] = creator.text

                date = meta_root.find('.//dc:date', meta_ns)
                if date is not None and date.text:
                    metadata['created'] = date.text

                keywords = meta_root.find('.//meta:keyword', meta_ns)
                if keywords is not None and keywords.text:
                    metadata['keywords'] = keywords.text
            except Exception as e:
                logger.warning(f"Failed to extract ODT metadata: {str(e)}")

        return text, metadata
    except Exception as e:
        logger.error(f"Error extracting text from ODT: {str(e)}")
        return "", {"error": str(e)}


def extract_text_from_file(file_content: bytes, file_extension: str) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text from a file based on its extension.

    Args:
        file_content: Binary file content
        file_extension: File extension (e.g., 'pdf', 'docx')

    Returns:
        Tuple of (extracted_text, metadata)
    """
    file_extension = file_extension.lower()

    extractors = {
        'pdf': extract_text_from_pdf,
        'docx': extract_text_from_docx,
        'doc': extract_text_from_docx,  # Will likely fail with .doc, but worth a try
        'txt': extract_text_from_txt,
        'rtf': extract_text_from_rtf,
        'odt': extract_text_from_odt
    }

    if file_extension not in extractors:
        logger.error(f"Unsupported file extension: {file_extension}")
        return "", {"error": f"Unsupported file extension: {file_extension}"}

    extractor = extractors[file_extension]
    return extractor(file_content)